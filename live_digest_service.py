#!/usr/bin/env python3
"""Monitor two Douyin rooms, record them, transcribe locally, and notify Feishu."""
from __future__ import annotations

import argparse
import asyncio
import hashlib
import json
import os
import sqlite3
import subprocess
import tempfile
import threading
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import requests


@dataclass
class Settings:
    recorder_root: Path
    output_dir: Path
    webhook: str
    app_id: str = ""
    app_secret: str = ""
    chat_id: str = ""
    recipient_open_ids: list[str] | None = None
    recipients: list[dict[str, str]] | None = None
    poll_seconds: int = 60
    segment_seconds: int = 900
    whisper_model: str = "small"
    whisper_language: str = "zh"
    transcription_mode: str = "server"
    proxy: str = ""
    cookie: str = ""
    bitable_app_token: str = ""
    account_table_id: str = ""
    record_table_id: str = ""
    config_poll_seconds: int = 60
    state_db: str = "./monitor_state.sqlite3"
    feishu_user_token_path: str = ""
    drive_root_folder_token: str = ""
    drive_platform_folder_name: str = "抖音"
    minutes_poll_seconds: int = 60
    minutes_timeout_seconds: int = 7200


@dataclass
class Account:
    account_id: str
    name: str
    room_url: str
    enabled: bool
    recipients: list[dict[str, str]]
    table_record_id: str = ""


class DeliveryLedger:
    """Durable per-session/per-recipient/per-message idempotency ledger."""

    def __init__(self, path: Path):
        self.db = sqlite3.connect(path, check_same_thread=False)
        self.lock = threading.Lock()
        with self.db:
            self.db.execute("CREATE TABLE IF NOT EXISTS deliveries (session_id TEXT, recipient_id TEXT, message_type TEXT, status TEXT, message_id TEXT, error TEXT, updated_at TEXT, PRIMARY KEY(session_id, recipient_id, message_type))")
            self.db.execute("CREATE TABLE IF NOT EXISTS sessions (account_id TEXT PRIMARY KEY, session_id TEXT, account_name TEXT, recipients TEXT, record_id TEXT, started_ms INTEGER, active INTEGER)")
            self.db.execute("CREATE TABLE IF NOT EXISTS service_flags (key TEXT PRIMARY KEY, value TEXT)")

    def claim(self, session_id: str, recipient_id: str, message_type: str) -> bool:
        with self.lock, self.db:
            row = self.db.execute("SELECT status FROM deliveries WHERE session_id=? AND recipient_id=? AND message_type=?", (session_id, recipient_id, message_type)).fetchone()
            if row and row[0] == "sent":
                return False
            self.db.execute("INSERT INTO deliveries VALUES(?,?,?,?,?,?,datetime('now')) ON CONFLICT(session_id,recipient_id,message_type) DO UPDATE SET status='sending', error=NULL, updated_at=datetime('now')", (session_id, recipient_id, message_type, "sending", "", ""))
            return True

    def finish(self, session_id: str, recipient_id: str, message_type: str, message_id: str = "", error: str = "") -> None:
        status = "sent" if not error else "failed"
        with self.lock, self.db:
            self.db.execute("UPDATE deliveries SET status=?, message_id=?, error=?, updated_at=datetime('now') WHERE session_id=? AND recipient_id=? AND message_type=?", (status, message_id, error, session_id, recipient_id, message_type))

    def active_session(self, account_id: str) -> dict[str, Any] | None:
        with self.lock:
            row = self.db.execute("SELECT session_id, account_name, recipients, record_id, started_ms FROM sessions WHERE account_id=? AND active=1", (account_id,)).fetchone()
        if not row:
            return None
        return {"session_id": row[0], "account_name": row[1], "recipients": json.loads(row[2]), "record_id": row[3], "started_ms": row[4]}

    def start_session(self, account_id: str, session_id: str, account_name: str, recipients: list[dict[str, str]], record_id: str, started_ms: int) -> bool:
        with self.lock, self.db:
            blocked = self.db.execute("SELECT value FROM service_flags WHERE key='deployment_pending'").fetchone()
            if blocked and blocked[0] == "1":
                return False
            self.db.execute("INSERT INTO sessions VALUES(?,?,?,?,?,?,1) ON CONFLICT(account_id) DO UPDATE SET session_id=excluded.session_id, account_name=excluded.account_name, recipients=excluded.recipients, record_id=excluded.record_id, started_ms=excluded.started_ms, active=1", (account_id, session_id, account_name, json.dumps(recipients, ensure_ascii=False), record_id, started_ms))
            return True

    def set_session_record_id(self, account_id: str, record_id: str) -> None:
        with self.lock, self.db:
            self.db.execute("UPDATE sessions SET record_id=? WHERE account_id=? AND active=1", (record_id, account_id))

    def set_deployment_pending(self, value: bool) -> None:
        with self.lock, self.db:
            self.db.execute("INSERT INTO service_flags VALUES('deployment_pending', ?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", ("1" if value else "0",))

    def end_session(self, account_id: str) -> None:
        with self.lock, self.db:
            self.db.execute("UPDATE sessions SET active=0 WHERE account_id=?", (account_id,))


_USER_TOKEN_LOCK = threading.Lock()


def recipient_targets(recipients: list[dict[str, str]] | None) -> list[tuple[str, str, str]]:
    result = []
    seen = set()
    for item in recipients or []:
        rid = item.get("id") or item.get("open_id")
        if not rid or rid in seen:
            continue
        seen.add(rid)
        result.append((item.get("id_type", "open_id"), rid, item.get("name", rid)))
    return result


def message_url(receive_type: str, session_id: str, recipient_id: str, message_type: str) -> str:
    """Use Feishu's request UUID so a transport retry cannot duplicate a message."""
    key = hashlib.sha256(f"{session_id}:{recipient_id}:{message_type}".encode()).hexdigest()
    return f"https://open.feishu.cn/open-apis/im/v1/messages?receive_id_type={receive_type}&uuid={key}"


def feishu_response_data(response: requests.Response) -> dict[str, Any]:
    """Fail on Feishu business errors, which can still use an HTTP 200 response."""
    response.raise_for_status()
    data = response.json()
    if data.get("code", 0) != 0:
        raise RuntimeError(f"Feishu API error {data.get('code')}: {data.get('msg', 'unknown error')}")
    return data


def feishu_text(webhook: str, text: str) -> None:
    if not webhook:
        print(text)
        return
    response = requests.post(webhook, json={"msg_type": "text", "content": {"text": text}}, timeout=20)
    feishu_response_data(response)


def send_text(settings: Settings, text: str, *, recipients: list[dict[str, str]] | None = None, session_id: str = "", message_type: str = "text", ledger: DeliveryLedger | None = None) -> None:
    """Send through the app bot when configured, otherwise use a webhook."""
    token = tenant_token(settings)
    targets = recipient_targets(recipients if recipients is not None else settings.recipients)
    if recipients is None:
        targets.extend(("open_id", value, value) for value in (settings.recipient_open_ids or []) if value)
    if settings.chat_id:
        targets.append(("chat_id", settings.chat_id, settings.chat_id))
    if token and targets:
        # A recipient may be present in both legacy and named settings. Keep
        # one delivery target per id so a single event cannot fan out twice.
        seen: set[tuple[str, str]] = set()
        for receive_type, receive_id, _ in targets:
            if (receive_type, receive_id) in seen:
                continue
            seen.add((receive_type, receive_id))
            if ledger and session_id and not ledger.claim(session_id, receive_id, message_type):
                continue
            try:
                endpoint = message_url(receive_type, session_id, receive_id, message_type) if session_id else f"https://open.feishu.cn/open-apis/im/v1/messages?receive_id_type={receive_type}"
                response = requests.post(endpoint, headers={"Authorization": f"Bearer {token}"}, json={"receive_id": receive_id, "msg_type": "text", "content": json.dumps({"text": text}, ensure_ascii=False)}, timeout=30)
                result = feishu_response_data(response)
                if ledger and session_id:
                    ledger.finish(session_id, receive_id, message_type, result.get("data", {}).get("message_id", ""))
            except Exception as exc:
                if ledger and session_id:
                    ledger.finish(session_id, receive_id, message_type, error=str(exc))
                raise
        return
    feishu_text(settings.webhook, text)


def tenant_token(settings: Settings) -> str | None:
    if not settings.app_id or not settings.app_secret:
        return None
    response = requests.post(
        "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
        json={"app_id": settings.app_id, "app_secret": settings.app_secret}, timeout=20,
    )
    data = feishu_response_data(response)
    return data.get("tenant_access_token")


def app_access_token(settings: Settings) -> str:
    if not settings.app_id or not settings.app_secret:
        raise RuntimeError("Feishu application credentials are not configured")
    response = requests.post(
        "https://open.feishu.cn/open-apis/auth/v3/app_access_token/internal",
        json={"app_id": settings.app_id, "app_secret": settings.app_secret}, timeout=20,
    )
    data = feishu_response_data(response)
    token = data.get("app_access_token")
    if not token:
        raise RuntimeError("Feishu did not return an app access token")
    return str(token)


def bitable_request(settings: Settings, method: str, path: str, body: dict[str, Any] | None = None) -> dict[str, Any]:
    token = tenant_token(settings)
    if not token or not settings.bitable_app_token:
        raise RuntimeError("Bitable credentials are not configured")
    response = requests.request(method, f"https://open.feishu.cn/open-apis{path}", headers={"Authorization": f"Bearer {token}"}, json=body, timeout=30)
    return feishu_response_data(response)


def user_token(settings: Settings) -> str:
    """Return the delegated token required by Feishu Minutes and Drive APIs."""
    if not settings.feishu_user_token_path:
        raise RuntimeError("Feishu user authorization token path is not configured")
    token_path = Path(settings.feishu_user_token_path)
    with _USER_TOKEN_LOCK:
        try:
            tokens = json.loads(token_path.read_text(encoding="utf-8"))
        except FileNotFoundError as exc:
            raise RuntimeError("Feishu user authorization is missing") from exc
        expires_at = float(tokens.get("expires_at", 0))
        if tokens.get("access_token") and expires_at > time.time() + 90:
            return str(tokens["access_token"])
        refresh_token = tokens.get("refresh_token")
        if not refresh_token or not settings.app_id or not settings.app_secret:
            raise RuntimeError("Feishu user authorization needs offline_access and a new authorization")
        response = requests.post(
            "https://open.feishu.cn/open-apis/authen/v1/oidc/refresh_access_token",
            json={"grant_type": "refresh_token", "refresh_token": refresh_token,
                  "app_access_token": app_access_token(settings)}, timeout=30,
        )
        data = feishu_response_data(response).get("data", {})
        access_token = data.get("access_token")
        if not access_token:
            raise RuntimeError("Feishu did not return a refreshed user token")
        tokens.update(data)
        tokens["expires_at"] = time.time() + int(data.get("expires_in", 7200))
        temporary = token_path.with_suffix(token_path.suffix + ".tmp")
        temporary.write_text(json.dumps(tokens, ensure_ascii=False), encoding="utf-8")
        temporary.chmod(0o600)
        temporary.replace(token_path)
        return str(access_token)


def user_feishu_request(settings: Settings, method: str, path: str, *, body: dict[str, Any] | None = None,
                        params: dict[str, Any] | None = None) -> dict[str, Any]:
    response = requests.request(
        method, f"https://open.feishu.cn/open-apis{path}", headers={"Authorization": f"Bearer {user_token(settings)}"},
        json=body, params=params, timeout=60,
    )
    return feishu_response_data(response)


def sync_accounts(settings: Settings) -> list[Account]:
    data = bitable_request(settings, "GET", f"/bitable/v1/apps/{settings.bitable_app_token}/tables/{settings.account_table_id}/records?page_size=500")
    accounts = []
    for item in data.get("data", {}).get("items", []):
        fields = item.get("fields", {})
        account_id = str(fields.get("抖音号", "")).strip()
        if not account_id:
            continue
        enabled = fields.get("监控开关") == "启用"
        url_value = fields.get("直播间链接")
        if isinstance(url_value, dict):
            url_value = url_value.get("link") or url_value.get("text")
        url = str(url_value or f"https://live.douyin.com/{account_id}")
        users = fields.get("监控接收人") or []
        recipients = [{"name": u.get("name", ""), "id_type": "open_id", "id": u.get("id", "")} for u in users if u.get("id")]
        accounts.append(Account(account_id, str(fields.get("监控账号", account_id)), url, enabled, recipients, item.get("record_id", item.get("id", ""))))
    return accounts


def update_account_state(settings: Settings, account: Account, *, status: str, started: int | None = None, ended: int | None = None, error: str | None = None) -> None:
    fields: dict[str, Any] = {"服务状态": status, "最后同步时间": int(time.time() * 1000)}
    if started:
        fields["最后开播时间"] = started
    if ended:
        fields["最后下播时间"] = ended
    if error:
        fields["服务状态"] = "异常"
    try:
        bitable_request(settings, "PUT", f"/bitable/v1/apps/{settings.bitable_app_token}/tables/{settings.account_table_id}/records/{account.table_record_id}", {"fields": fields})
    except Exception as exc:
        print(f"Account state update failed for {account.account_id}: {exc}", flush=True)


def create_live_record(settings: Settings, account: Account, session_id: str, title: str, started: int, recipients: list[dict[str, str]]) -> str:
    label = f"【{account.name}】{datetime.fromtimestamp(started / 1000).strftime('%Y%m%d_%H%M')}-进行中"
    fields = {"直播记录": label, "账号名称": account.name, "抖音号": account.account_id, "直播标题": title, "开播时间": started, "录制状态": "录制中", "转写状态": "待下载", "推送状态": "待推送", "任务 ID": session_id}
    data = bitable_request(settings, "POST", f"/bitable/v1/apps/{settings.bitable_app_token}/tables/{settings.record_table_id}/records", {"fields": fields})
    return data.get("data", {}).get("record", {}).get("record_id", "")


def update_live_record(settings: Settings, record_id: str, fields: dict[str, Any]) -> None:
    if not record_id:
        return
    try:
        bitable_request(settings, "PUT", f"/bitable/v1/apps/{settings.bitable_app_token}/tables/{settings.record_table_id}/records/{record_id}", {"fields": fields})
    except Exception as exc:
        print(f"Live record update failed: {exc}", flush=True)


def upload_bitable_attachment(settings: Settings, path: Path) -> str | None:
    """Upload an artifact for a Bitable attachment field and return its token."""
    token = tenant_token(settings)
    if not token or not settings.bitable_app_token or not path.exists():
        return None
    with path.open("rb") as handle:
        response = requests.post(
            "https://open.feishu.cn/open-apis/drive/v1/medias/upload_all",
            headers={"Authorization": f"Bearer {token}"},
            data={"file_name": path.name, "parent_type": "bitable_file", "parent_node": settings.bitable_app_token, "size": str(path.stat().st_size)},
            files={"file": (path.name, handle)}, timeout=300,
        )
    return feishu_response_data(response).get("data", {}).get("file_token")


def attach_session_artifacts(settings: Settings, record_id: str, screenshot: Path, *, minute_url: str = "",
                             transcript_url: str = "") -> None:
    """Attach the screenshot and write the durable Feishu links for a session."""
    fields: dict[str, Any] = {}
    image_token = upload_bitable_attachment(settings, screenshot)
    if image_token:
        fields["截图"] = [{"file_token": image_token, "name": screenshot.name}]
    if minute_url:
        fields["智能纪要链接"] = {"link": minute_url, "text": "打开智能纪要"}
    if transcript_url:
        fields["文字记录链接"] = {"link": transcript_url, "text": "打开文字记录"}
    if fields:
        update_live_record(settings, record_id, fields)


def concat_segments(segments: list[Path], output: Path) -> Path:
    """Join recorder segments into the one complete MP4 that gets archived."""
    if not segments:
        raise RuntimeError("No recording segments to merge")
    output.parent.mkdir(parents=True, exist_ok=True)
    if len(segments) == 1:
        if segments[0].resolve() != output.resolve():
            subprocess.run(["ffmpeg", "-y", "-hide_banner", "-loglevel", "error", "-i", str(segments[0]),
                            "-c", "copy", "-movflags", "+faststart", str(output)], check=True)
        return output
    with tempfile.NamedTemporaryFile("w", encoding="utf-8", suffix=".ffconcat", delete=False) as handle:
        concat_file = Path(handle.name)
        for segment in segments:
            # ffconcat requires single quotes to be escaped in file paths.
            handle.write(f"file '{str(segment.resolve()).replace(chr(39), chr(39) + chr(92) + chr(39) + chr(39))}'\\n")
    try:
        fast_command = ["ffmpeg", "-y", "-hide_banner", "-loglevel", "error", "-f", "concat", "-safe", "0",
                        "-i", str(concat_file), "-c", "copy", "-movflags", "+faststart", str(output)]
        copied = subprocess.run(fast_command, check=False)
        if copied.returncode != 0:
            subprocess.run(["ffmpeg", "-y", "-hide_banner", "-loglevel", "error", "-f", "concat", "-safe", "0",
                            "-i", str(concat_file), "-c:v", "libx264", "-c:a", "aac", "-movflags", "+faststart",
                            str(output)], check=True)
    finally:
        concat_file.unlink(missing_ok=True)
    return output


def drive_list(settings: Settings, folder_token: str) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    page_token = ""
    while True:
        params: dict[str, Any] = {"folder_token": folder_token, "page_size": 200}
        if page_token:
            params["page_token"] = page_token
        data = user_feishu_request(settings, "GET", "/drive/v1/files", params=params).get("data", {})
        items.extend(data.get("files", []))
        if not data.get("has_more"):
            return items
        page_token = data.get("next_page_token", "")
        if not page_token:
            return items


def ensure_drive_folder(settings: Settings, parent_token: str, name: str) -> str:
    for item in drive_list(settings, parent_token):
        if item.get("name") == name and item.get("type") == "folder":
            return str(item.get("token"))
    data = user_feishu_request(settings, "POST", "/drive/v1/files/create_folder",
                               body={"name": name, "folder_token": parent_token})
    token = data.get("data", {}).get("token")
    if not token:
        raise RuntimeError(f"Could not create Feishu folder: {name}")
    return str(token)


def session_drive_folder(settings: Settings, account_name: str) -> str:
    if not settings.drive_root_folder_token:
        raise RuntimeError("Feishu Drive root folder token is not configured")
    platform = ensure_drive_folder(settings, settings.drive_root_folder_token, settings.drive_platform_folder_name)
    account = ensure_drive_folder(settings, platform, account_name)
    return ensure_drive_folder(settings, account, "直播录像")


def drive_file_url(file_token: str) -> str:
    return f"https://shenyidushu.feishu.cn/drive/file/{file_token}"


def upload_drive_file(settings: Settings, path: Path, folder_token: str) -> str:
    """Upload to a Drive folder with the multipart API, suitable for full MP4 files."""
    if not path.is_file():
        raise RuntimeError(f"Artifact does not exist: {path}")
    prepared = user_feishu_request(settings, "POST", "/drive/v1/files/upload_prepare", body={
        "file_name": path.name, "parent_type": "explorer", "parent_node": folder_token, "size": path.stat().st_size,
    }).get("data", {})
    upload_id = prepared.get("upload_id")
    block_size = int(prepared.get("block_size", 0))
    block_num = int(prepared.get("block_num", 0))
    if not upload_id or block_size <= 0 or block_num <= 0:
        raise RuntimeError("Feishu did not prepare the file upload")
    token = user_token(settings)
    with path.open("rb") as source:
        for sequence in range(block_num):
            block = source.read(block_size)
            if not block:
                raise RuntimeError("Recording ended before all upload blocks were read")
            response = requests.post(
                f"https://open.feishu.cn/open-apis/drive/v1/files/{upload_id}/upload_part",
                headers={"Authorization": f"Bearer {token}"}, params={"seq": sequence},
                files={"file": (path.name, block)}, timeout=300,
            )
            feishu_response_data(response)
    finished = user_feishu_request(settings, "POST", "/drive/v1/files/upload_finish",
                                   body={"upload_id": upload_id, "block_num": block_num}).get("data", {})
    file_token = finished.get("file_token")
    if not file_token:
        raise RuntimeError("Feishu did not return the uploaded file token")
    return str(file_token)


def upload_minutes(settings: Settings, file_token: str) -> tuple[str, str]:
    data = user_feishu_request(settings, "POST", "/minutes/v1/minutes/upload", body={"file_token": file_token}).get("data", {})
    token = data.get("minute_token")
    url = data.get("minute_url")
    if not token or not url:
        raise RuntimeError("Feishu did not create a Minutes record for this video")
    return str(token), str(url)


def wait_for_transcript(settings: Settings, minute_token: str) -> str:
    deadline = time.monotonic() + settings.minutes_timeout_seconds
    while time.monotonic() < deadline:
        response = requests.get(
            f"https://open.feishu.cn/open-apis/minutes/v1/minutes/{minute_token}/transcript",
            headers={"Authorization": f"Bearer {user_token(settings)}"}, timeout=60,
        )
        if response.status_code == 200:
            return response.text.strip()
        try:
            data = response.json()
        except ValueError:
            response.raise_for_status()
            raise RuntimeError("Feishu Minutes returned an unreadable transcript response")
        if data.get("code") != 2091003:
            feishu_response_data(response)
        time.sleep(settings.minutes_poll_seconds)
    raise RuntimeError("Feishu Minutes transcription timed out")


def create_transcript_docx(path: Path, account_name: str, session_id: str, transcript: str) -> Path:
    from docx import Document

    document = Document()
    document.add_heading(f"{account_name} 直播逐字稿", level=1)
    document.add_paragraph(f"直播开始时间：{artifact_timestamp(session_id)}")
    for paragraph in transcript.splitlines():
        document.add_paragraph(paragraph)
    document.save(path)
    return path


def upload_file(settings: Settings, path: Path) -> str | None:
    """Upload a message file. Feishu limits this endpoint to 30 MB."""
    token = tenant_token(settings)
    if not token:
        return None
    with path.open("rb") as video:
        response = requests.post(
            "https://open.feishu.cn/open-apis/im/v1/files",
            headers={"Authorization": f"Bearer {token}"},
            # `file_size` is not accepted by this endpoint and turns an
            # otherwise valid upload into a 400 business error.
            data={"file_type": "stream", "file_name": path.name},
            files={"file": (path.name, video)}, timeout=300,
        )
    return feishu_response_data(response).get("data", {}).get("file_key")


def upload_image(settings: Settings, path: Path) -> str | None:
    """Upload a JPG screenshot for an inline Feishu image message."""
    token = tenant_token(settings)
    if not token:
        return None
    with path.open("rb") as image:
        response = requests.post(
            "https://open.feishu.cn/open-apis/im/v1/images",
            headers={"Authorization": f"Bearer {token}"},
            data={"image_type": "message"},
            files={"image": (path.name, image, "image/jpeg")}, timeout=60,
        )
    return feishu_response_data(response).get("data", {}).get("image_key")


def feishu_file(settings: Settings, file_key: str, text: str, *, recipients: list[dict[str, str]] | None = None, session_id: str = "", ledger: DeliveryLedger | None = None) -> None:
    token = tenant_token(settings)
    targets = recipient_targets(recipients if recipients is not None else settings.recipients)
    if recipients is None:
        targets.extend(("open_id", value, value) for value in (settings.recipient_open_ids or []) if value)
    if settings.chat_id:
        targets.append(("chat_id", settings.chat_id, settings.chat_id))
    if not token or not targets:
        send_text(settings, text)
        return
    seen: set[tuple[str, str]] = set()
    for receive_type, receive_id, _ in targets:
        if (receive_type, receive_id) in seen:
            continue
        seen.add((receive_type, receive_id))
        if ledger and session_id and not ledger.claim(session_id, receive_id, "transcript"):
            continue
        try:
            response = requests.post(
            message_url(receive_type, session_id, receive_id, "transcript") if session_id else f"https://open.feishu.cn/open-apis/im/v1/messages?receive_id_type={receive_type}",
            headers={"Authorization": f"Bearer {token}"},
            json={"receive_id": receive_id, "msg_type": "file", "content": json.dumps({"file_key": file_key})},
            timeout=30,
        )
            result = feishu_response_data(response)
            if ledger and session_id:
                ledger.finish(session_id, receive_id, "transcript", result.get("data", {}).get("message_id", ""))
        except Exception as exc:
            if ledger and session_id:
                ledger.finish(session_id, receive_id, "transcript", error=str(exc))
            raise
    # The attachment itself is the transcript delivery. Do not call
    # send_text here: that used to create a duplicate notification per user.


def feishu_image(settings: Settings, image_key: str, *, recipients: list[dict[str, str]] | None = None, session_id: str = "", ledger: DeliveryLedger | None = None) -> None:
    token = tenant_token(settings)
    targets = recipient_targets(recipients if recipients is not None else settings.recipients)
    if recipients is None:
        targets.extend(("open_id", value, value) for value in (settings.recipient_open_ids or []) if value)
    if settings.chat_id:
        targets.append(("chat_id", settings.chat_id, settings.chat_id))
    if not token or not targets:
        return
    seen: set[tuple[str, str]] = set()
    for receive_type, receive_id, _ in targets:
        if (receive_type, receive_id) in seen:
            continue
        seen.add((receive_type, receive_id))
        if ledger and session_id and not ledger.claim(session_id, receive_id, "screenshot"):
            continue
        try:
            response = requests.post(
            message_url(receive_type, session_id, receive_id, "screenshot") if session_id else f"https://open.feishu.cn/open-apis/im/v1/messages?receive_id_type={receive_type}",
            headers={"Authorization": f"Bearer {token}"},
            json={"receive_id": receive_id, "msg_type": "image",
                  "content": json.dumps({"image_key": image_key})},
            timeout=30,
        )
            result = feishu_response_data(response)
            if ledger and session_id:
                ledger.finish(session_id, receive_id, "screenshot", result.get("data", {}).get("message_id", ""))
        except Exception as exc:
            if ledger and session_id:
                ledger.finish(session_id, receive_id, "screenshot", error=str(exc))
            raise


def artifact_timestamp(session_id: str) -> str:
    """Make the session start time readable while keeping it filename-safe."""
    try:
        return datetime.strptime(session_id, "%Y%m%d_%H%M%S").strftime("%Y-%m-%d_%H-%M-%S")
    except ValueError:
        return session_id


def artifact_path(directory: Path, kind: str, account_id: str, session_id: str, suffix: str) -> Path:
    return directory / f"{kind}-{account_id}-{artifact_timestamp(session_id)}{suffix}"


def capture_screenshot(video: Path, screenshot: Path) -> None:
    """Capture a readable frame near the beginning of the first recording segment."""
    subprocess.run([
        "ffmpeg", "-y", "-hide_banner", "-loglevel", "error", "-ss", "5", "-i", str(video),
        "-frames:v", "1", "-vf", "scale='min(1280,iw)':-2", "-q:v", "3", str(screenshot),
    ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)


def publish_finished_session(
    settings: Settings, *, first_segment: Path, transcript: Path, account_id: str,
    session_id: str, anchor: str, title: str, url: str, transcript_length: int,
    recipients: list[dict[str, str]], ledger: DeliveryLedger,
) -> None:
    """Send the requested screenshot and named transcript, never the oversized MP4."""
    screenshot = artifact_path(transcript.parent, "直播截图", anchor or account_id, session_id, ".jpg")
    capture_screenshot(first_segment, screenshot)
    image_key = upload_image(settings, screenshot)
    if image_key:
        feishu_image(settings, image_key, recipients=recipients, session_id=session_id, ledger=ledger)
    else:
        send_text(settings, f"【直播截图】{anchor}\n本地文件：{screenshot}", recipients=recipients, session_id=session_id, message_type="screenshot", ledger=ledger)

    transcript_key = upload_file(settings, transcript)
    caption = f"【下播，完整逐字稿】{anchor}\n{title}\n{url}\n共 {transcript_length} 字。"
    if transcript_key:
        feishu_file(settings, transcript_key, caption, recipients=recipients, session_id=session_id, ledger=ledger)
    else:
        send_text(settings, f"{caption}\n本地文件：{transcript}", recipients=recipients, session_id=session_id, message_type="transcript", ledger=ledger)


def complete_with_feishu_minutes(
    settings: Settings, *, room_dir: Path, segments: list[Path], account_name: str, session_id: str,
    record_id: str, title: str, url: str, recipients: list[dict[str, str]], ledger: DeliveryLedger,
) -> None:
    """Create one complete video, archive it, transcribe it in Minutes, then publish once."""
    complete_video = artifact_path(room_dir, "直播视频", account_name, session_id, "_00.mp4")
    concat_segments(segments, complete_video)
    screenshot = artifact_path(room_dir, "直播截图", account_name, session_id, ".jpg")
    capture_screenshot(complete_video, screenshot)
    update_live_record(settings, record_id, {"录制状态": "已完成", "转写状态": "转写中"})

    archive_folder = session_drive_folder(settings, account_name)
    video_token = upload_drive_file(settings, complete_video, archive_folder)
    minute_token, minute_url = upload_minutes(settings, video_token)
    transcript_text = wait_for_transcript(settings, minute_token)
    transcript_docx = artifact_path(room_dir, "直播逐字稿", account_name, session_id, ".docx")
    create_transcript_docx(transcript_docx, account_name, session_id, transcript_text)
    transcript_token = upload_drive_file(settings, transcript_docx, archive_folder)
    transcript_url = drive_file_url(transcript_token)

    publish_finished_session(
        settings, first_segment=complete_video, transcript=transcript_docx, account_id=account_name,
        session_id=session_id, anchor=account_name, title=title, url=url, transcript_length=len(transcript_text),
        recipients=recipients, ledger=ledger,
    )
    attach_session_artifacts(settings, record_id, screenshot, minute_url=minute_url, transcript_url=transcript_url)
    update_live_record(settings, record_id, {
        "录制状态": "已完成", "转写状态": "已完成", "推送状态": "已推送",
        "推送时间": int(time.time() * 1000), "失败原因": "",
    })


def transcribe(path: Path, settings: Settings) -> str:
    """Use openai-whisper's CLI; output is kept beside the media for recovery."""
    out_dir = path.parent / "transcripts"
    out_dir.mkdir(exist_ok=True)
    stem = out_dir / path.stem
    txt = stem.with_suffix(".txt")
    if txt.exists():
        return txt.read_text(encoding="utf-8")
    wav = path.with_suffix(".wav")
    subprocess.run(["ffmpeg", "-y", "-i", str(path), "-vn", "-ar", "16000", "-ac", "1", str(wav)],
                   check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    subprocess.run(["whisper", str(wav), "--model", settings.whisper_model, "--language",
                    settings.whisper_language, "--output_format", "txt", "--output_dir", str(out_dir)], check=True)
    wav.unlink(missing_ok=True)
    return txt.read_text(encoding="utf-8") if txt.exists() else ""


def stream_info(settings: Settings, url: str) -> dict[str, Any]:
    import sys
    sys.path.insert(0, str(settings.recorder_root))
    from src import spider, stream  # type: ignore
    data = asyncio.run(spider.get_douyin_web_stream_data(url, proxy_addr=settings.proxy or None,
                                                         cookies=settings.cookie or None))
    return asyncio.run(stream.get_douyin_stream_url(data, "OD", settings.proxy or None))


def start_recorder(
    settings: Settings, room_dir: Path, account_id: str, session_stamp: str, record_url: str,
) -> subprocess.Popen[bytes]:
    """Start or resume a segmented recording without overwriting existing segments."""
    video_base = artifact_path(room_dir, "直播视频", account_id, session_stamp, "")
    existing = [segment for segment in room_dir.glob(f"{video_base.name}_*.mp4")
                if len(segment.stem.rsplit("_", 1)[-1]) == 3 and segment.stem.rsplit("_", 1)[-1].isdigit()]
    indices = []
    for segment in existing:
        try:
            indices.append(int(segment.stem.rsplit("_", 1)[-1]))
        except ValueError:
            continue
    next_index = max(indices, default=-1) + 1
    command = [
        "ffmpeg", "-hide_banner", "-loglevel", "warning", "-i", record_url,
        "-c", "copy", "-f", "segment", "-segment_time", str(settings.segment_seconds),
        "-segment_format", "mp4", "-segment_start_number", str(next_index),
        "-reset_timestamps", "1", "-movflags", "+faststart", f"{video_base}_%03d.mp4",
    ]
    return subprocess.Popen(command)


def session_segments(room_dir: Path, account_name: str, session_id: str) -> list[Path]:
    """Return only temporary recorder segments, never the final merged MP4."""
    prefix = f"直播视频-{account_name}-{artifact_timestamp(session_id)}_"
    return [segment for segment in sorted(room_dir.glob(f"{prefix}*.mp4"))
            if segment.stat().st_size >= 1024
            and len(segment.stem.rsplit("_", 1)[-1]) == 3
            and segment.stem.rsplit("_", 1)[-1].isdigit()]


def run_room(settings: Settings, account_id: str, registry: dict[str, Account], registry_lock: threading.Lock, ledger: DeliveryLedger) -> None:
    room_dir = settings.output_dir / account_id
    room_dir.mkdir(parents=True, exist_ok=True)
    active = False
    process: subprocess.Popen[bytes] | None = None
    session_stamp: str | None = None
    session_snapshot: dict[str, Any] | None = ledger.active_session(account_id)
    if session_snapshot:
        active = True
        session_stamp = session_snapshot["session_id"]
    print(f"Monitoring {account_id} every {settings.poll_seconds}s", flush=True)
    while True:
        try:
            with registry_lock:
                account = registry.get(account_id)
            if not account:
                time.sleep(settings.poll_seconds)
                continue
            url = account.room_url
            if not account.enabled and not active:
                update_account_state(settings, account, status="未使用")
                time.sleep(settings.poll_seconds)
                continue
            info = stream_info(settings, url)
            live = bool(info.get("is_live") and info.get("record_url"))
            if live and not active:
                active = True
                session_stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                started_ms = int(time.time() * 1000)
                if not ledger.start_session(account_id, session_stamp, account.name, account.recipients, "", started_ms):
                    active, session_stamp = False, None
                    time.sleep(settings.poll_seconds)
                    continue
                process = start_recorder(settings, room_dir, account.name, session_stamp, info["record_url"])
                record_id = create_live_record(settings, account, session_stamp, info.get("title", ""), started_ms, account.recipients)
                ledger.set_session_record_id(account_id, record_id)
                session_snapshot = {"account_name": account.name, "recipients": account.recipients, "record_id": record_id, "started_ms": started_ms}
                update_account_state(settings, account, status="正常使用", started=started_ms)
                send_text(settings, f"【开播】{account.name}\n{info.get('title', '')}\n{url}", recipients=account.recipients, session_id=session_stamp, message_type="live_start", ledger=ledger)
                print(f"Live started: {url} ({session_stamp})", flush=True)
            elif live and active and process and process.poll() is not None:
                exit_code = process.returncode
                process = start_recorder(
                    settings, room_dir, account.name, session_stamp or datetime.now().strftime("%Y%m%d_%H%M%S"),
                    info["record_url"],
                )
                print(f"Recorder restarted: {url}; previous exit code {exit_code}", flush=True)
            elif live and active and process is None:
                snap_name = (session_snapshot or {}).get("account_name", account.name)
                process = start_recorder(settings, room_dir, snap_name, session_stamp or datetime.now().strftime("%Y%m%d_%H%M%S"), info["record_url"])
            if active and not live:
                if process:
                    process.terminate()
                    process.wait(timeout=30)
                snap_name = (session_snapshot or {}).get("account_name", account.name)
                snap_recipients = (session_snapshot or {}).get("recipients", account.recipients)
                record_id = (session_snapshot or {}).get("record_id", "")
                started_ms = (session_snapshot or {}).get("started_ms", int(time.time() * 1000))
                segments = session_segments(room_dir, snap_name, session_stamp or "unknown")
                if settings.transcription_mode == "local_pull":
                    manifest = room_dir / f"{session_stamp}_pending_transcription.json"
                    manifest.write_text(json.dumps({
                        "session_id": session_stamp,
                        "account_id": account_id,
                        "account_name": snap_name,
                        "recipient_snapshot": snap_recipients,
                        "record_id": record_id,
                        "url": url,
                        "anchor_name": account.name,
                        "title": info.get("title", ""),
                        "segments": [segment.name for segment in segments],
                        "created_at": datetime.now().isoformat(timespec="seconds"),
                    }, ensure_ascii=False, indent=2), encoding="utf-8")
                    ended_ms = int(time.time() * 1000)
                    update_live_record(settings, record_id, {"直播记录": f"【{snap_name}】{datetime.fromtimestamp(started_ms / 1000).strftime('%Y%m%d_%H%M')}-{datetime.fromtimestamp(ended_ms / 1000).strftime('%H%M')}", "下播时间": ended_ms, "直播时长（分钟）": max(0, int((ended_ms - started_ms) / 60000)), "录制状态": "已完成" if segments else "录制失败", "转写状态": "待下载", "推送状态": "待推送"})
                    ledger.end_session(account_id)
                    update_account_state(settings, account, status="正常使用", ended=ended_ms)
                    print(f"Local transcription queued: {manifest}", flush=True)
                    active, process, session_stamp = False, None, None
                    time.sleep(settings.poll_seconds)
                    continue
                if settings.transcription_mode == "feishu_minutes":
                    ended_ms = int(time.time() * 1000)
                    update_live_record(settings, record_id, {
                        "直播记录": f"【{snap_name}】{datetime.fromtimestamp(started_ms / 1000).strftime('%Y%m%d_%H%M')}-{datetime.fromtimestamp(ended_ms / 1000).strftime('%H%M')}",
                        "下播时间": ended_ms, "直播时长（分钟）": max(0, int((ended_ms - started_ms) / 60000)),
                        "录制状态": "已完成" if segments else "录制失败", "转写状态": "转写中" if segments else "转写失败",
                    })
                    if not segments:
                        update_live_record(settings, record_id, {"失败原因": "下播后未发现有效录像分段", "推送状态": "未推送"})
                    else:
                        try:
                            complete_with_feishu_minutes(
                                settings, room_dir=room_dir, segments=segments, account_name=snap_name,
                                session_id=session_stamp or "unknown", record_id=record_id,
                                title=info.get("title", ""), url=url, recipients=snap_recipients, ledger=ledger,
                            )
                        except Exception as exc:
                            update_live_record(settings, record_id, {
                                "转写状态": "转写失败", "推送状态": "未推送", "失败原因": str(exc)[:1000],
                            })
                            raise
                    ledger.end_session(account_id)
                    update_account_state(settings, account, status="正常使用", ended=ended_ms)
                    print(f"Feishu Minutes processing finished: {url}; {len(segments)} segments", flush=True)
                    active, process, session_stamp, session_snapshot = False, None, None, None
                    time.sleep(settings.poll_seconds)
                    continue
                transcripts = []
                for segment in segments:
                    transcripts.append(transcribe(segment, settings))
                full = "\n\n".join(text for text in transcripts if text)
                full_path = artifact_path(room_dir, "直播逐字稿", account.name, session_stamp or "unknown", ".txt")
                full_path.write_text(full, encoding="utf-8")
                first_segment = segments[0] if segments else None
                if first_segment:
                    publish_finished_session(
                        settings, first_segment=first_segment, transcript=full_path, account_id=account_id,
                        session_id=session_stamp or "unknown", anchor=account.name,
                        title=info.get("title", ""), url=url, transcript_length=len(full), recipients=snap_recipients, ledger=ledger,
                    )
                    attach_session_artifacts(
                        settings, record_id,
                        artifact_path(room_dir, "直播截图", snap_name, session_stamp or "unknown", ".jpg"),
                    )
                update_live_record(settings, record_id, {"录制状态": "已完成", "转写状态": "已完成", "推送状态": "已推送", "推送时间": int(time.time() * 1000)})
                print(f"Live finished: {url}; {len(segments)} segments, {len(full)} chars", flush=True)
                active, process, session_stamp, session_snapshot = False, None, None, None
            time.sleep(settings.poll_seconds)
        except Exception as exc:
            print(f"{account_id}: {exc}", flush=True)
            time.sleep(settings.poll_seconds)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--config", default="live_digest.json")
    args = parser.parse_args()
    cfg = json.loads(Path(args.config).read_text(encoding="utf-8"))
    settings = Settings(recorder_root=Path(cfg["recorder_root"]), output_dir=Path(cfg.get("output_dir", "./recordings")),
                        webhook=cfg.get("feishu_webhook", ""), app_id=cfg.get("feishu_app_id", ""),
                        app_secret=cfg.get("feishu_app_secret", ""), chat_id=cfg.get("feishu_chat_id", ""),
                        recipient_open_ids=cfg.get("feishu_open_ids", []),
                        recipients=cfg.get("feishu_recipients", []),
                        poll_seconds=int(cfg.get("poll_seconds", 60)),
                        segment_seconds=int(cfg.get("segment_seconds", 900)),
                        whisper_model=cfg.get("whisper_model", "small"),
                        whisper_language=cfg.get("whisper_language", "zh"), proxy=cfg.get("proxy", ""),
                        transcription_mode=cfg.get("transcription_mode", "server"),
                        cookie=cfg.get("douyin_cookie", ""), bitable_app_token=cfg.get("bitable_app_token", ""),
                        account_table_id=cfg.get("account_table_id", ""), record_table_id=cfg.get("record_table_id", ""),
                        config_poll_seconds=int(cfg.get("config_poll_seconds", 60)), state_db=cfg.get("state_db", "./monitor_state.sqlite3"),
                        feishu_user_token_path=cfg.get("feishu_user_token_path", ""),
                        drive_root_folder_token=cfg.get("drive_root_folder_token", ""),
                        drive_platform_folder_name=cfg.get("drive_platform_folder_name", "抖音"),
                        minutes_poll_seconds=int(cfg.get("minutes_poll_seconds", 60)),
                        minutes_timeout_seconds=int(cfg.get("minutes_timeout_seconds", 7200)))
    settings.output_dir.mkdir(parents=True, exist_ok=True)
    registry: dict[str, Account] = {}
    registry_lock = threading.Lock()
    ledger = DeliveryLedger(Path(settings.state_db))
    if settings.bitable_app_token and settings.account_table_id:
        for account in sync_accounts(settings):
            registry[account.account_id] = account
    else:
        for url in cfg.get("rooms", []):
            account_id = url.rstrip("/").split("/")[-1]
            registry[account_id] = Account(account_id, account_id, url, True, settings.recipients or [])

    started: set[str] = set()
    threads: list[threading.Thread] = []
    stop = threading.Event()

    def ensure_threads() -> None:
        while not stop.is_set():
            if settings.bitable_app_token and settings.account_table_id:
                try:
                    accounts = sync_accounts(settings)
                    with registry_lock:
                        registry.clear()
                        registry.update({a.account_id: a for a in accounts})
                    # This column represents whether the monitoring service is
                    # available for an account, not its transient live state.
                    for account in accounts:
                        update_account_state(
                            settings, account,
                            status="正常使用" if account.enabled else "未使用",
                        )
                except Exception as exc:
                    print(f"Account config sync failed: {exc}", flush=True)
            with registry_lock:
                ids = list(registry)
            for account_id in ids:
                if account_id not in started:
                    started.add(account_id)
                    thread = threading.Thread(target=run_room, args=(settings, account_id, registry, registry_lock, ledger), daemon=True)
                    thread.start()
                    threads.append(thread)
            stop.wait(settings.config_poll_seconds)

    manager = threading.Thread(target=ensure_threads, daemon=True)
    manager.start()
    while True:
        time.sleep(3600)


if __name__ == "__main__":
    main()
